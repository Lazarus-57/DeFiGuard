import docx
from docx.shared import Pt, Inches
from pathlib import Path

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    # Make headings a bit larger and darker
    for run in heading.runs:
        run.font.color.rgb = docx.shared.RGBColor(38, 70, 83)
        if level == 0:
            run.font.size = Pt(20)
        elif level == 1:
            run.font.size = Pt(16)

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    return p

def main():
    doc = docx.Document()
    
    # Title
    add_heading(doc, 'DeFIGuard - 25k Dataset & Model Evaluation Report', level=0)
    
    # Section 1
    add_heading(doc, '1. The 25k Dataset: Composition', level=1)
    add_paragraph(doc, 
        "In this phase of the project, we scaled our dataset to approximately 25,000 transactions. "
        "Our goal is to build an Anti-Money Laundering (AML) detection pipeline that can accurately "
        "find suspicious transactions hiding among legitimate ones."
    )
    add_paragraph(doc, 
        "To train and test the model, we injected two complex, real-world money laundering patterns into the data:\n"
        " - Peel Chains: A long sequence of transactions where a small amount is 'peeled' off at each step, "
        "acting like a chain of drops to slowly cash out funds.\n"
        " - Smurfing: A fan-out and fan-in pattern where a single source wallet distributes funds to multiple "
        "'mule' wallets, which then forward the money to a central collector wallet. This hides the large transfer."
    )
    
    add_paragraph(doc, 
        "Dataset Breakdown:\n"
        " - Total Transactions: 26,854\n"
        " - Normal (Legitimate): 25,000 (93.1%)\n"
        " - Peel Chains: 1,050 (3.9%)\n"
        " - Smurfing: 804 (3.0%)"
    )

    # Insert Image
    img_path = Path("reports/25k/Pipeline Overview/composition_pie_25k.png")
    if img_path.exists():
        doc.add_picture(str(img_path), width=Inches(4.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Section 2
    add_heading(doc, '2. Model Evaluation and Layers', level=1)
    add_paragraph(doc, 
        "How do we know if our model is doing a good job? Since suspicious transactions only make up about 7% of the data, "
        "traditional 'Accuracy' can be misleading (a model that guesses 'Normal' every time would still be 93% accurate!). "
        "Instead, we use PR-AUC (Precision-Recall Area Under Curve), which heavily penalizes false alarms and rewards "
        "correctly catching the rare suspicious events."
    )
    add_paragraph(doc, "We tested our models in three layers of increasing complexity:")
    
    doc.add_heading('Layer 1: Traditional Baselines', level=2)
    add_paragraph(doc, 
        "We started with standard machine learning models using basic transaction features (like amounts and simple network degrees).\n"
        " - Random Forest: 0.8603 PR-AUC\n"
        " - Logistic Regression: 0.9015 PR-AUC\n"
        " - XGBoost: 0.9311 PR-AUC\n"
        "XGBoost was the clear winner here, setting a strong baseline."
    )

    doc.add_heading('Layer 2: Adding Graph Context (GraphSAGE)', level=2)
    add_paragraph(doc, 
        "Next, we gave the model 'Graph Context'. We used an advanced Graph Neural Network (GraphSAGE) to map the shape "
        "and structure of the wallets' transaction history. Did giving XGBoost this structural awareness help?\n"
        " - XGBoost Baseline: 0.9250 PR-AUC\n"
        " - XGBoost + GraphSAGE: 0.9351 PR-AUC (+0.0101)\n"
        "Yes! The graph features helped the model better recognize the structural signatures of smurfing and peel chains."
    )

    doc.add_heading('Layer 3: Adding Temporal Intelligence (NTS)', level=2)
    add_paragraph(doc, 
        "Finally, we tested 'Temporal Intelligence'. Money laundering is rarely instantaneous; it involves deliberate delays "
        "and time-spreads. We engineered Network Time Spread (NTS) features to capture the rhythm and timing of transactions.\n"
        " - XGBoost Baseline: 0.9250 PR-AUC\n"
        " - XGBoost + NTS: 0.9365 PR-AUC (+0.0114)\n"
        "This yielded the biggest jump. The model learned that the timing of transactions is just as suspicious as the structure."
    )

    # Section 3
    add_heading(doc, '3. The Final Winner', level=1)
    add_paragraph(doc, 
        "The final winner is XGBoost augmented with Temporal Intelligence (NTS), achieving the highest PR-AUC score of 0.9365."
    )
    add_paragraph(doc, 
        "Why did this win?\n"
        "While Graph Neural Networks (GraphSAGE) successfully captured the spatial layout of laundering operations, the time-based "
        "features proved to be an even stronger signal. Both peel chains and smurfing rely on specific sequence timings "
        "to move funds securely. By giving XGBoost awareness of these temporal rhythms, the model became highly effective "
        "at distinguishing complex money laundering from normal DeFi activity."
    )

    out_path = Path("reports/25k/DeFIGuard_25k_Report.docx")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

if __name__ == '__main__':
    main()
