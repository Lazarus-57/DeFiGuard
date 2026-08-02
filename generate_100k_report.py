import os
from pathlib import Path
from docx import Document
from fpdf import FPDF

def create_docx(output_path):
    doc = Document()
    doc.add_heading('DeFIGuard: 100k Pipeline Evaluation Report', 0)
    
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        "This report details the findings from the 100k dataset evaluation pipeline. We sequentially "
        "injected three synthetic laundering patterns (Peel Chains, Smurfing Clusters, and Circular Rings) "
        "into a dataset of 100,000 legitimate Ethereum transactions. Our goal was to evaluate how traditional "
        "machine learning (XGBoost), structural Graph Neural Networks (GraphSAGE), and Temporal Intelligence (NTS) "
        "perform under increasing laundering complexity."
    )
    
    doc.add_heading('1. Addressing the Temporal Data Leakage', level=1)
    doc.add_paragraph(
        "During our initial review, we discovered a data leakage bug in our legacy pipelines (such as the 25k run). "
        "The synthetic transactions were assigned completely random timestamps, maximizing their Node Temporal Spread (NTS) "
        "artificially. In this 100k pipeline run, we fixed this by strictly enforcing realistic chronological offsets "
        "(e.g., 1-5 minute delays between hops, and tightly batched fan-outs for smurfing). This provides a mathematically "
        "sound and realistic evaluation."
    )

    doc.add_heading('2. Dataset Overview', level=1)
    doc.add_paragraph("Step 1: Peel Chains Only (4.0% positive rate, 104,144 total transactions)")
    doc.add_paragraph("Step 2: Peel Chains + Smurfing (6.7% positive rate, 107,220 total transactions)")
    doc.add_paragraph("Step 3: All Patterns (Peel + Smurf + Circular) (9.4% positive rate, 110,362 total transactions)")
    
    doc.add_heading('3. Key Findings by Step', level=1)
    
    doc.add_heading('Step 1: Peel Chains Only', level=2)
    doc.add_paragraph("Baseline PR-AUC: 0.1238")
    doc.add_paragraph(
        "GraphSAGE caused a performance drop (-0.0230) due to oversmoothing on the linear chains. "
        "NTS Temporal features also showed a slight drop (-0.0239) because realistic peel chains look temporally "
        "similar to regular sequential DeFi usage."
    )
    
    doc.add_heading('Step 2: Peel Chains + Smurfing', level=2)
    doc.add_paragraph("Baseline PR-AUC: 0.1563")
    doc.add_paragraph(
        "GraphSAGE provided a strong lift (+0.0539) because the fan-out/fan-in topology forms distinct structural clusters. "
        "However, NTS Temporal features provided a massive lift (+0.1064) because the coordinated, batched nature of smurfing "
        "creates an unmistakable temporal burst signature."
    )
    
    doc.add_heading('Step 3: All Patterns (Peel + Smurf + Circular)', level=2)
    doc.add_paragraph("Baseline PR-AUC: 0.2405")
    doc.add_paragraph(
        "GraphSAGE failed on the circular loops (-0.0315), as neighborhood aggregations get confused when funds loop back "
        "to the source. NTS Temporal features provided a slight lift (+0.0033) as the sequential nature of circular rings "
        "is slightly harder to hide temporally than structurally."
    )

    doc.add_heading('4. Conclusion', level=1)
    doc.add_paragraph(
        "The 100k pipeline demonstrates that Graph Neural Networks (GraphSAGE) are highly sensitive to graph topology—excelling "
        "at cluster detection but failing on linear or cyclic structures. Conversely, Temporal Intelligence (NTS) is incredibly "
        "effective at detecting coordinated laundering behavior (like smurfing fan-outs) that require tight temporal synchronization."
    )
    
    doc.save(output_path)


class PDFReport(FPDF):
    def header(self):
        self.set_font('Helvetica', 'B', 15)
        self.cell(0, 10, 'DeFIGuard: 100k Pipeline Evaluation Report', 0, 1, 'C')
        self.ln(10)

    def chapter_title(self, title):
        self.set_font('Helvetica', 'B', 12)
        self.cell(0, 10, title, 0, 1, 'L')
        self.ln(4)

    def chapter_body(self, body):
        self.set_font('Helvetica', '', 11)
        self.multi_cell(0, 6, body)
        self.ln()

def create_pdf(output_path):
    pdf = PDFReport()
    pdf.add_page()
    
    pdf.chapter_title('Executive Summary')
    pdf.chapter_body(
        "This report details the findings from the 100k dataset evaluation pipeline. We sequentially "
        "injected three synthetic laundering patterns (Peel Chains, Smurfing Clusters, and Circular Rings) "
        "into a dataset of 100,000 legitimate Ethereum transactions. Our goal was to evaluate how traditional "
        "machine learning (XGBoost), structural Graph Neural Networks (GraphSAGE), and Temporal Intelligence (NTS) "
        "perform under increasing laundering complexity."
    )
    
    pdf.chapter_title('1. Addressing the Temporal Data Leakage')
    pdf.chapter_body(
        "During our initial review, we discovered a data leakage bug in our legacy pipelines (such as the 25k run). "
        "The synthetic transactions were assigned completely random timestamps, maximizing their Node Temporal Spread (NTS) "
        "artificially. In this 100k pipeline run, we fixed this by strictly enforcing realistic chronological offsets "
        "(e.g., 1-5 minute delays between hops, and tightly batched fan-outs for smurfing). This provides a mathematically "
        "sound and realistic evaluation."
    )
    
    pdf.chapter_title('2. Dataset Overview')
    pdf.chapter_body(
        "Step 1: Peel Chains Only (4.0% positive rate, 104,144 total transactions)\n"
        "Step 2: Peel Chains + Smurfing (6.7% positive rate, 107,220 total transactions)\n"
        "Step 3: All Patterns (Peel + Smurf + Circular) (9.4% positive rate, 110,362 total transactions)"
    )

    pdf.chapter_title('3. Key Findings by Step')
    
    pdf.set_font('Helvetica', 'B', 11)
    pdf.cell(0, 6, 'Step 1: Peel Chains Only', 0, 1, 'L')
    pdf.chapter_body(
        "Baseline PR-AUC: 0.1238\n"
        "GraphSAGE caused a performance drop (-0.0230) due to oversmoothing on the linear chains. "
        "NTS Temporal features also showed a slight drop (-0.0239) because realistic peel chains look temporally "
        "similar to regular sequential DeFi usage."
    )
    
    pdf.set_font('Helvetica', 'B', 11)
    pdf.cell(0, 6, 'Step 2: Peel Chains + Smurfing', 0, 1, 'L')
    pdf.chapter_body(
        "Baseline PR-AUC: 0.1563\n"
        "GraphSAGE provided a strong lift (+0.0539) because the fan-out/fan-in topology forms distinct structural clusters. "
        "However, NTS Temporal features provided a massive lift (+0.1064) because the coordinated, batched nature of smurfing "
        "creates an unmistakable temporal burst signature."
    )
    
    pdf.set_font('Helvetica', 'B', 11)
    pdf.cell(0, 6, 'Step 3: All Patterns (Peel + Smurf + Circular)', 0, 1, 'L')
    pdf.chapter_body(
        "Baseline PR-AUC: 0.2405\n"
        "GraphSAGE failed on the circular loops (-0.0315), as neighborhood aggregations get confused when funds loop back "
        "to the source. NTS Temporal features provided a slight lift (+0.0033) as the sequential nature of circular rings "
        "is slightly harder to hide temporally than structurally."
    )

    pdf.chapter_title('4. Conclusion')
    pdf.chapter_body(
        "The 100k pipeline demonstrates that Graph Neural Networks (GraphSAGE) are highly sensitive to graph topology-excelling "
        "at cluster detection but failing on linear or cyclic structures. Conversely, Temporal Intelligence (NTS) is incredibly "
        "effective at detecting coordinated laundering behavior (like smurfing fan-outs) that require tight temporal synchronization."
    )
    
    pdf.output(output_path)


if __name__ == "__main__":
    reports_dir = Path("reports/100k")
    reports_dir.mkdir(parents=True, exist_ok=True)
    
    docx_path = reports_dir / "DeFIGuard_100k_Report.docx"
    pdf_path = reports_dir / "DeFIGuard_100k_Report.pdf"
    
    create_docx(str(docx_path))
    print(f"Created {docx_path}")
    
    create_pdf(str(pdf_path))
    print(f"Created {pdf_path}")
